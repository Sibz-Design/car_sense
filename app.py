from flask import Flask, render_template, jsonify, request, send_file
import requests
import json
from datetime import datetime, timedelta, timezone
import os
import re
from textblob import TextBlob
import logging
from jinja2.exceptions import TemplateNotFound
from dotenv import load_dotenv
from docx import Document
from collections import Counter
import io
from bs4 import BeautifulSoup

load_dotenv()

app = Flask(__name__)

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# YouTube API Configuration
YOUTUBE_API_KEY_1 = os.getenv("API_KEY_1")
YOUTUBE_API_KEY_2 = os.getenv("API_KEY_2")
CHANNEL_ID = "UCB-mfYAd3oJLEkoMxjRAxbg"

class YouTubeCommentsService:
    def __init__(self):
        self.api_keys = [YOUTUBE_API_KEY_1, YOUTUBE_API_KEY_2]
        self.channel_id = CHANNEL_ID
        self.current_api_key_index = 0
    
    def get_current_api_key(self):
        """Get the current API key"""
        return self.api_keys[self.current_api_key_index]
    
    def switch_api_key(self):
        """Switch to the next API key if rate limit or error occurs"""
        self.current_api_key_index = (self.current_api_key_index + 1) % len(self.api_keys)
        logger.info(f"Switched to API key {self.current_api_key_index + 1}")
        return self.get_current_api_key()
    
    def analyze_sentiment(self, text):
        """Analyze sentiment of text using TextBlob"""
        try:
            cleaned_text = re.sub(r'<[^>]+>', '', text)
            cleaned_text = re.sub(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', '', cleaned_text)
            cleaned_text = re.sub(r'[^\w\s]', '', cleaned_text)
            
            if not cleaned_text.strip():
                return 'neutral'
            
            blob = TextBlob(cleaned_text)
            polarity = blob.sentiment.polarity
            
            if polarity > 0.1:
                return 'positive'
            elif polarity < -0.1:
                return 'negative'
            else:
                return 'neutral'
        except Exception as e:
            logger.error(f"Error analyzing sentiment: {e}")
            return 'neutral'
    
    def get_latest_videos(self, max_results=50):
        """Get latest videos from the YouTube channel"""
        url = "https://www.googleapis.com/youtube/v3/search"
        published_after = (datetime.now(timezone.utc) - timedelta(days=30)).replace(microsecond=0).isoformat()
        params = {
            'key': self.get_current_api_key(),
            'channelId': self.channel_id,
            'part': 'snippet,id',
            'order': 'date',
            'maxResults': min(max_results, 50),
            'type': 'video',
            'publishedAfter': published_after
        }
        
        for attempt in range(len(self.api_keys)):
            try:
                logger.info(f"Fetching videos with params: {params}")
                response = requests.get(url, params=params, timeout=30)
                response.raise_for_status()
                data = response.json()
                
                if 'error' in data:
                    logger.error(f"YouTube API error: {data['error']}")
                    if data['error'].get('code') == 403:
                        params['key'] = self.switch_api_key()
                        continue
                    else:
                        raise Exception(data['error']['message'])
                
                videos = []
                for item in data.get('items', []):
                    if 'videoId' in item.get('id', {}):
                        videos.append({
                            'videoId': item['id']['videoId'],
                            'title': item['snippet']['title'],
                            'publishedAt': item['snippet']['publishedAt'],
                            'description': item['snippet'].get('description', '')[:200],
                            'thumbnail': item['snippet']['thumbnails'].get('default', {}).get('url', '')
                        })
                
                logger.info(f"Retrieved {len(videos)} videos")
                return videos
                
            except requests.exceptions.RequestException as e:
                logger.error(f"Error fetching videos (attempt {attempt} + 1): {e}")
                if isinstance(e, requests.exceptions.HTTPError) and e.response.status_code == 403:
                    params['key'] = self.switch_api_key()
                    continue
                return []
            except Exception as e:
                logger.error(f"Unexpected error fetching videos: {e}")
                return []
        
        logger.error("All API keys failed to fetch videos")
        return []
    
    def get_comments_for_video(self, video_id, max_results=50):
        """Get comments for a specific video"""
        url = "https://www.googleapis.com/youtube/v3/commentThreads"
        params = {
            'key': self.get_current_api_key(),
            'part': 'snippet',
            'videoId': video_id,
            'maxResults': min(max_results, 100),
            'order': 'time'
        }
        
        for attempt in range(len(self.api_keys)):
            try:
                logger.info(f"Fetching comments for video {video_id} with params: {params}")
                response = requests.get(url, params=params, timeout=30)
                response.raise_for_status()
                data = response.json()
                
                if 'error' in data:
                    logger.error(f"YouTube API error for video {video_id}: {data['error']}")
                    if data['error'].get('code') == 403:
                        params['key'] = self.switch_api_key()
                        continue
                    else:
                        raise Exception(data['error']['message'])
                
                comments = []
                for item in data.get('items', []):
                    try:
                        comment_data = item['snippet']['topLevelComment']['snippet']
                        comment_text = comment_data['textDisplay']
                        sentiment = self.analyze_sentiment(comment_text)
                        
                        comments.append({
                            'author': comment_data['authorDisplayName'],
                            'comment': comment_text[:500],
                            'date': comment_data['publishedAt'],
                            'likeCount': comment_data.get('likeCount', 0),
                            'sentiment': sentiment,
                            'authorProfileImageUrl': comment_data.get('authorProfileImageUrl', ''),
                            'videoId': video_id
                        })
                    except KeyError as e:
                        logger.warning(f"Missing key in comment data: {e}")
                        continue
                
                logger.info(f"Retrieved {len(comments)} comments for video {video_id}")
                return comments
                
            except requests.exceptions.RequestException as e:
                logger.error(f"Error fetching comments for video {video_id} (attempt {attempt} + 1): {e}")
                if isinstance(e, requests.exceptions.HTTPError) and e.response.status_code == 403:
                    params['key'] = self.switch_api_key()
                    continue
                return []
            except Exception as e:
                logger.error(f"Unexpected error fetching comments for video {video_id}: {e}")
                return []
        
        logger.error(f"All API keys failed to fetch comments for video {video_id}")
        return []
    
    def get_video_details_by_url(self, video_url, max_comments=50):
        """Get video details and comments based on a YouTube or Bing search URL"""
        try:
            # Check if URL is a Bing search results page
            if 'bing.com/videos' in video_url:
                logger.info(f"Processing Bing search URL: {video_url}")
                headers = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                }
                response = requests.get(video_url, headers=headers, timeout=30)
                response.raise_for_status()
                soup = BeautifulSoup(response.text, 'html.parser')
                
                # Find the first YouTube video link
                youtube_url = None
                for link in soup.find_all('a', href=True):
                    href = link['href']
                    if 'youtube.com/watch?v=' in href or 'youtu.be/' in href:
                        youtube_url = href
                        break
                
                if not youtube_url:
                    raise ValueError("No YouTube video found in Bing search results")
                
                # Ensure full URL
                if youtube_url.startswith('/'):
                    youtube_url = 'https://www.bing.com' + youtube_url
                elif not youtube_url.startswith('http'):
                    youtube_url = 'https://' + youtube_url
                
                logger.info(f"Extracted YouTube URL: {youtube_url}")
                video_url = youtube_url
        
            # Extract video ID from YouTube URL
            video_id = None
            if 'youtube.com/watch?v=' in video_url:
                video_id = video_url.split('v=')[1].split('&')[0]
            elif 'youtu.be/' in video_url:
                video_id = video_url.split('youtu.be/')[1].split('?')[0]
            
            if not video_id:
                raise ValueError("Invalid YouTube URL")
            
            # Fetch video details
            url = "https://www.googleapis.com/youtube/v3/videos"
            params = {
                'key': self.get_current_api_key(),
                'part': 'snippet',
                'id': video_id
            }
            
            for attempt in range(len(self.api_keys)):
                try:
                    response = requests.get(url, params=params, timeout=30)
                    response.raise_for_status()
                    data = response.json()
                    
                    if 'error' in data:
                        logger.error(f"YouTube API error for video {video_id}: {data['error']}")
                        if data['error'].get('code') == 403:
                            params['key'] = self.switch_api_key()
                            continue
                        else:
                            raise Exception(data['error']['message'])
                    
                    if not data.get('items'):
                        raise ValueError("Video not found")
                    
                    video_data = data['items'][0]['snippet']
                    comments = self.get_comments_for_video(video_id, max_comments)
                    
                    return {
                        'videoId': video_id,
                        'title': video_data['title'],
                        'publishedAt': video_data['publishedAt'],
                        'description': video_data.get('description', '')[:200],
                        'thumbnail': video_data['thumbnails'].get('default', {}).get('url', ''),
                        'comments': comments,
                        'commentCount': len(comments)
                    }
                
                except requests.exceptions.RequestException as e:
                    logger.error(f"Error fetching video details for {video_id} (attempt {attempt} + 1): {e}")
                    if isinstance(e, requests.exceptions.HTTPError) and e.response.status_code == 403:
                        params['key'] = self.switch_api_key()
                        continue
                    raise
                except Exception as e:
                    logger.error(f"Unexpected error fetching video details for {video_id}: {e}")
                    raise
            
            logger.error(f"All API keys failed to fetch video details for {video_id}")
            raise Exception("Failed to fetch video details")
        
        except Exception as e:
            logger.error(f"Error in get_video_details_by_url: {e}")
            raise
    
    def get_all_comments_data(self, max_videos=10, max_comments_per_video=50):
        """Get all comments data for analysis"""
        try:
            videos = self.get_latest_videos(max_videos)
            all_comments = []
            video_comment_counts = {}
            videos_with_comments = []
            
            for i, video in enumerate(videos[:max_videos]):
                logger.info(f"Processing video {i+1}/{max_videos}: {video['title'][:50]}...")
                comments = self.get_comments_for_video(video['videoId'], max_comments_per_video)
                
                if comments:
                    video_title_short = video['title'][:30] + ('...' if len(video['title']) > 30 else '')
                    video_comment_counts[video_title_short] = len(comments)
                    videos_with_comments.append({
                        'title': video['title'],
                        'videoId': video['videoId'],
                        'publishedAt': video['publishedAt'],
                        'description': video.get('description', ''),
                        'thumbnail': video.get('thumbnail', ''),
                        'comments': comments,
                        'commentCount': len(comments)
                    })
                    all_comments.extend(comments)
            
            sentiment_counts = {'positive': 0, 'negative': 0, 'neutral': 0}
            for comment in all_comments:
                sentiment_counts[comment['sentiment']] += 1
            
            total_likes = sum(comment['likeCount'] for comment in all_comments)
            avg_likes_per_comment = total_likes / len(all_comments) if all_comments else 0
            
            logger.info(f"Analysis complete: {len(all_comments)} comments from {len(videos_with_comments)} videos")
            
            return {
                'total_comments': len(all_comments),
                'video_comment_counts': video_comment_counts,
                'comments': all_comments,
                'videos_with_comments': videos_with_comments,
                'total_videos': len(videos_with_comments),
                'sentiment_counts': sentiment_counts,
                'total_likes': total_likes,
                'avg_likes_per_comment': round(avg_likes_per_comment, 2),
                'processed_at': datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Error in get_all_comments_data: {e}")
            return {
                'total_comments': 0,
                'video_comment_counts': {},
                'comments': [],
                'videos_with_comments': [],
                'total_videos': 0,
                'sentiment_counts': {'positive': 0, 'negative': 0, 'neutral': 0},
                'total_likes': 0,
                'avg_likes_per_comment': 0,
                'error': str(e)
            }

# Initialize the service
youtube_service = YouTubeCommentsService()

# AI Analysis Function
def generate_ai_analysis(video_data, sentiment_type='negative'):
    """Generate detailed AI analysis for specified sentiment comments of a video"""
    try:
        comments = [c for c in video_data['comments'] if c['sentiment'] == sentiment_type]
        if not comments:
            return {
                'overview': f"No {sentiment_type} comments found for the video '{video_data['title']}'. This suggests that the video has been well-received in terms of {sentiment_type} feedback. Consider maintaining or enhancing the elements that contributed to this response in future content.",
                'comments_by_video': [{'title': video_data['title'], 'comments': []}],
                'themes': [],
                'recommendations': []
            }
        
        # Group comments by video (single video in this case)
        comments_by_video = [{
            'title': video_data['title'],
            'comments': sorted(comments, key=lambda x: x['likeCount'], reverse=True)[:10]
        }]
        
        # Extract themes from comments
        words = []
        for comment in comments:
            cleaned_text = re.sub(r'[^\w\s]', '', comment['comment'].lower())
            words.extend(cleaned_text.split())
        
        common_words = Counter(words).most_common(20)
        themes = [word for word, count in common_words if len(word) > 3 and word not in ['this', 'that', 'with', 'from', 'have', 'they', 'video', 'like', 'dont']]
        
        # Generate recommendations based on sentiment type
        recommendations = []
        for theme in themes[:5]:
            if sentiment_type == 'negative':
                recommendations.append(
                    f"Address concerns related to '{theme}': Review comments mentioning '{theme}' to identify specific issues. For example, if '{theme}' relates to content quality, "
                    "improve production aspects like editing, audio clarity, or visuals. If '{theme}' involves topic relevance, align future content with audience interests through surveys or engagement metrics."
                )
            else:  # positive
                recommendations.append(
                    f"Reinforce positive feedback on '{theme}': Continue emphasizing '{theme}' in future content, as it resonates well with viewers. For example, if '{theme}' relates to informative content, "
                    "create more in-depth tutorials or reviews to maintain audience engagement."
                )
        
        recommendations.extend([
            f"Engage with {sentiment_type} commenters: Respond to the top {sentiment_type} comments listed below to {'address concerns' if sentiment_type == 'negative' else 'acknowledge praise'}. "
            f"This can {'improve trust' if sentiment_type == 'negative' else 'strengthen loyalty'} and show commitment to audience feedback.",
            f"Track {sentiment_type} sentiment trends: Monitor {'negative' if sentiment_type == 'negative' else 'positive'} feedback in future videos to {'identify persistent issues' if sentiment_type == 'negative' else 'sustain positive engagement'}. Adjust content strategy accordingly.",
            "Experiment with content formats: Based on feedback, try new formats like live streams or Q&A sessions to address audience preferences or reinforce successful elements."
        ])
        
        overview = (
            f"Analyzed {len(comments)} {sentiment_type} comments for the video '{video_data['title']}' from the provided URL. "
            f"The analysis below lists the top 10 {sentiment_type} comments, identifies key themes driving this sentiment, and provides actionable recommendations to "
            f"{'improve content and reduce negative sentiment' if sentiment_type == 'negative' else 'maintain and enhance positive engagement'}. "
            "Implementing these suggestions can help optimize viewer satisfaction and engagement."
        )
        
        return {
            'overview': overview,
            'comments_by_video': comments_by_video,
            'themes': themes[:5],
            'recommendations': recommendations
        }
    except Exception as e:
        logger.error(f"Error in generate_ai_analysis: {e}")
        return {
            'overview': f"Error generating {sentiment_type} analysis due to an unexpected issue.",
            'comments_by_video': [{'title': video_data.get('title', 'Unknown Video'), 'comments': []}],
            'themes': [],
            'recommendations': [],
            'error': str(e)
        }

@app.route('/')
def dashboard():
    """Main dashboard page"""
    try:
        return render_template('index.html')
    except TemplateNotFound:
        logger.error("Template 'index.html' not found in templates directory")
        return jsonify({'error': 'Template index.html not found'}), 500

@app.route('/api/chart-data')
def get_chart_data():
    """Get data formatted for charts"""
    max_videos = request.args.get('max_videos', 10, type=int)
    max_comments = request.args.get('max_comments', 50, type=int)
    
    max_videos = min(max(max_videos, 1), 20)
    max_comments = min(max(max_comments, 10), 100)
    
    try:
        data = youtube_service.get_all_comments_data(max_videos, max_comments)
        
        video_counts = data['video_comment_counts']
        sorted_videos = sorted(video_counts.items(), key=lambda x: x[1], reverse=True)[:10]
        
        pie_data = {
            'labels': [video[0] for video in sorted_videos],
            'values': [video[1] for video in sorted_videos],
            'colors': ['#FF0000', '#00b894', '#fdcb6e', '#54A0FF', '#5F27CD', 
                      '#FF9FF3', '#96CEB4', '#FECA57', '#45B7D1', '#FF9F43'][:len(sorted_videos)]
        }
        
        comments_by_date = {}
        for comment in data['comments']:
            date = comment['date'][:10]
            comments_by_date[date] = comments_by_date.get(date, 0) + 1
        
        sorted_dates = sorted(comments_by_date.items())[-30:]
        
        bar_data = {
            'labels': [item[0] for item in sorted_dates],
            'values': [item[1] for item in sorted_dates]
        }
        
        sentiment_by_date = {}
        for comment in data['comments']:
            date = comment['date'][:10]
            if date not in sentiment_by_date:
                sentiment_by_date[date] = {'positive': 0, 'negative': 0, 'neutral': 0}
            sentiment_by_date[date][comment['sentiment']] += 1
        
        sentiment_trend = {
            'dates': sorted(sentiment_by_date.keys())[-14:],
            'positive': [],
            'negative': [],
            'neutral': []
        }
        
        for date in sentiment_trend['dates']:
            day_data = sentiment_by_date.get(date, {'positive': 0, 'negative': 0, 'neutral': 0})
            sentiment_trend['positive'].append(day_data['positive'])
            sentiment_trend['negative'].append(day_data['negative'])
            sentiment_trend['neutral'].append(day_data['neutral'])
        
        return jsonify({
            'pie_chart': pie_data,
            'bar_chart': bar_data,
            'sentiment_trend': sentiment_trend,
            'summary': {
                'total_comments': data['total_comments'],
                'total_videos': data['total_videos'],
                'sentiment_counts': data['sentiment_counts'],
                'total_likes': data.get('total_likes', 0),
                'avg_likes_per_comment': data.get('avg_likes_per_comment', 0)
            }
        })
    except Exception as e:
        logger.error(f"Error in get_chart_data: {e}")
        return jsonify({
            'error': str(e)
        }), 500

@app.route('/api/sentiment-data')
def get_sentiment_data():
    """Get detailed sentiment data with comments"""
    max_videos = request.args.get('max_videos', 5, type=int)
    max_comments = request.args.get('max_comments', 20, type=int)
    
    max_videos = min(max(max_videos, 1), 10)
    max_comments = min(max(max_comments, 10), 50)
    
    try:
        data = youtube_service.get_all_comments_data(max_videos, max_comments)
        
        sample_comments = {'positive': [], 'negative': [], 'neutral': []}
        for comment in data['comments']:
            sentiment = comment['sentiment']
            if len(sample_comments[sentiment]) < 10:
                sample_comments[sentiment].append({
                    'author': comment['author'],
                    'comment': comment['comment'][:200],
                    'likeCount': comment['likeCount'],
                    'date': comment['date'],
                    'videoId': comment['videoId']
                })
        
        return jsonify({
            'videos_with_comments': data['videos_with_comments'],
            'sentiment_summary': data['sentiment_counts'],
            'sample_comments': sample_comments,
            'total_comments': data['total_comments'],
            'total_videos': data['total_videos'],
            'total_likes': data.get('total_likes', 0),
            'avg_likes_per_comment': data.get('avg_likes_per_comment', 0)
        })
    except Exception as e:
        logger.error(f"Error in get_sentiment_data: {e}")
        return jsonify({
            'error': str(e)
        }), 500

@app.route('/api/video-details/<video_id>')
def get_video_details(video_id):
    """Get detailed information about a specific video"""
    max_comments = request.args.get('max_comments', 100, type=int)
    
    try:
        comments = youtube_service.get_comments_for_video(video_id, max_comments)
        
        sentiment_counts = {'positive': 0, 'negative': 0, 'neutral': 0}
        for comment in comments:
            sentiment_counts[comment['sentiment']] += 1
        
        return jsonify({
            'video_id': video_id,
            'comments': comments,
            'comment_count': len(comments),
            'sentiment_counts': sentiment_counts,
            'total_likes': sum(comment['likeCount'] for comment in comments)
        })
    except Exception as e:
        logger.error(f"Error getting video details for {video_id}: {e}")
        return jsonify({
            'error': str(e)
        }), 500

@app.route('/api/ai-analysis')
def get_ai_analysis():
    """Get AI analysis for a specific video's comments based on URL and sentiment type"""
    video_url = request.args.get('video_url', '')
    sentiment_type = request.args.get('sentiment_type', 'negative')
    
    if not video_url:
        return jsonify({'error': 'Video URL is required'}), 400
    if sentiment_type not in ['positive', 'negative']:
        return jsonify({'error': 'Invalid sentiment type'}), 400
    
    try:
        video_data = youtube_service.get_video_details_by_url(video_url, max_comments=50)
        analysis = generate_ai_analysis(video_data, sentiment_type)
        return jsonify(analysis)
    except Exception as e:
        logger.error(f"Error in get_ai_analysis: {e}")
        return jsonify({
            'error': str(e)
        }), 500

@app.route('/api/export-data')
def export_data():
    """Export AI analysis as Word document"""
    format_type = request.args.get('format', 'DOCX').upper()
    video_url = request.args.get('video_url', '')
    sentiment_type = request.args.get('sentiment_type', 'negative')
    
    if format_type != 'DOCX':
        return jsonify({'error': 'Only DOCX format is supported'}), 400
    if not video_url:
        return jsonify({'error': 'Video URL is required for export'}), 400
    if sentiment_type not in ['positive', 'negative']:
        return jsonify({'error': 'Invalid sentiment type'}), 400
    
    try:
        video_data = youtube_service.get_video_details_by_url(video_url, max_comments=50)
        analysis = generate_ai_analysis(video_data, sentiment_type)
        
        doc = Document()
        doc.add_heading('YouTube Comments AI Analysis Report', 0)
        
        doc.add_heading('Overview', level=1)
        doc.add_paragraph(analysis['overview'])
        
        doc.add_heading(f'Top {sentiment_type.capitalize()} Comments', level=1)
        for video in analysis['comments_by_video']:
            doc.add_heading(video['title'], level=2)
            for comment in video['comments']:
                doc.add_paragraph(f"{comment['author']}: {comment['comment']} (Likes: {comment['likeCount']})", style='List Bullet')
        
        doc.add_heading('Key Themes', level=1)
        for theme in analysis['themes']:
            doc.add_paragraph(theme, style='List Bullet')
        
        doc.add_heading('Recommendations for Improvement', level=1)
        for recommendation in analysis['recommendations']:
            doc.add_paragraph(recommendation, style='List Bullet')
        
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return send_file(
            doc_buffer,
            as_attachment=True,
            download_name=f"ai_analysis_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    
    except Exception as e:
        logger.error(f"Error generating export: {e}")
        return jsonify({
            'error': str(e)
        }), 500

@app.errorhandler(404)
def not_found(error):
    try:
        return render_template('404.html'), 404
    except TemplateNotFound:
        return jsonify({'error': 'Page not found'}), 404

@app.errorhandler(500)
def internal_error(error):
    try:
        return render_template('500.html'), 500
    except TemplateNotFound:
        return jsonify({'error': 'Internal server error'}), 500

if __name__ == '__main__':
    os.makedirs('templates', exist_ok=True)
    os.makedirs('static/css', exist_ok=True)
    os.makedirs('static/js', exist_ok=True)
    
    app.run(debug=True, host='0.0.0.0', port=5000)
